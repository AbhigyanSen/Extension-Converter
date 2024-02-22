from flask import Flask, render_template, request, send_file
import os
from werkzeug.utils import secure_filename
from io import BytesIO
import PyPDF2
from docx import Document
from fpdf import FPDF

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Limit file size to 16 MB

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_document_type(extension):
    document_types = {
        'xls': 'Document: Spreadsheet',
        'xlsx': 'Document: Spreadsheet',
        'doc': 'Document: Document File',
        'docx': 'Document: Document File',
        'pdf': 'Document: Portable Document',
        'txt': 'Document: Text Document',
    }
    return document_types.get(extension.lower(), 'Document: Unknown Type')

def convert_file(input_path, output_dir, output_format):
    filename_without_extension, _ = os.path.splitext(os.path.basename(input_path))
    
    if output_format == 'docx':
        docx_output_path = os.path.join(output_dir, f"{filename_without_extension}.docx")
        txt_to_docx(input_path, docx_output_path)
    elif output_format == 'pdf':
        pdf_output_path = os.path.join(output_dir, f"{filename_without_extension}.pdf")
        if output_format == 'docx':
            docx_to_pdf(input_path, pdf_output_path)
        elif output_format == 'txt':
            txt_to_pdf(input_path, pdf_output_path)
    elif output_format == 'txt':
        txt_output_path = os.path.join(output_dir, f"{filename_without_extension}.txt")
        if output_format == 'docx':
            docx_to_txt(input_path, txt_output_path)
        elif output_format == 'pdf':
            pdf_to_txt(input_path, txt_output_path)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('index.html', error='No file part')

        file = request.files['file']

        if file.filename == '':
            return render_template('index.html', error='No selected file')

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            _, file_extension = os.path.splitext(filename)
            document_type = get_document_type(file_extension[1:])

            output_formats = ['docx', 'pdf', 'txt']
            output_formats.remove(file_extension[1:])

            return render_template('convert.html', document_type=document_type, output_formats=output_formats, filename=filename)
    return render_template('index.html')

@app.route('/convert/<filename>', methods=['POST'])
def convert(filename):
    output_format = request.form.get('output_format')
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if output_format and output_format.lower() in ALLOWED_EXTENSIONS:
        convert_file(file_path, app.config['UPLOAD_FOLDER'], output_format.lower())
        return render_template('download.html', filename=filename, output_format=output_format.lower())
    else:
        return render_template('index.html', error='Invalid output format')

@app.route('/download/<filename>/<output_format>')
def download(filename, output_format):
    converted_filename = f"{os.path.splitext(filename)[0]}.{output_format.lower()}"
    converted_filepath = os.path.join(app.config['UPLOAD_FOLDER'], converted_filename)
    return send_file(converted_filepath, as_attachment=True)

def replace_problematic_characters(text):
    problematic_chars = ['\u2019', '\u2013']  # Add other problematic characters if needed
    for char in problematic_chars:
        text = text.replace(char, '?')
    return text

def pdf_to_txt(input_path, output_dir):
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''.join(reader.pages[page_num].extract_text() for page_num in range(len(reader.pages)))

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.txt')
    with open(output_path, 'w') as txt_file:
        txt_file.write(text)

def pdf_to_docx(input_path, output_dir):
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        doc = Document()
        for page_num in range(len(reader.pages)):
            text = reader.pages[page_num].extract_text()
            doc.add_paragraph(text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.docx')
    doc.save(output_path)

def txt_to_pdf(input_path, output_dir):
    with open(input_path, 'r', encoding='utf-8') as file:
        text = replace_problematic_characters(file.read())

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, txt=text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.pdf')
    pdf.output(output_path)

def txt_to_docx(input_path, output_dir):
    with open(input_path, 'r') as file:
        text = file.read()

    doc = Document()
    doc.add_paragraph(text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.docx')
    doc.save(output_path)

def docx_to_pdf(input_path, output_dir):
    doc = Document(input_path)
    text = replace_problematic_characters(para.text + '\n' for para in doc.paragraphs)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, txt=text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.pdf')
    pdf.output(output_path)

def docx_to_txt(input_path, output_dir):
    doc = Document(input_path)
    text = ''.join(para.text + '\n' for para in doc.paragraphs)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.txt')
    with open(output_path, 'w') as txt_file:
        txt_file.write(text)

if __name__ == '__main__':
    app.run(debug=True)