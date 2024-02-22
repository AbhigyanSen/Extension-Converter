import os
import flask
import PyPDF2
from docx import Document
from fpdf import FPDF
import uuid

app = flask.Flask(__name__)
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def upload_file(file):
    if file and allowed_file(file.filename):
        filename = str(uuid.uuid4()) + '.' + file.filename.rsplit('.', 1)[1].lower()
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        return filename
    return None

def convert_file(input_file, output_format):
    try:
        output_filename = str(uuid.uuid4()) + '.' + output_format
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        if input_file.endswith('.pdf'):
            if output_format == 'txt':
                with open(input_file, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ''.join(reader.pages[page_num].extract_text() for page_num in range(len(reader.pages)))
                    with open(output_path, 'w', encoding='utf-8') as txt_file:
                        txt_file.write(text)
            elif output_format == 'docx':
                with open(input_file, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    doc = Document()
                    for page_num in range(len(reader.pages)):
                        text = reader.pages[page_num].extract_text()
                        doc.add_paragraph(text)
                    doc.save(output_path)
            else:
                return None  # Same format conversion not supported

        elif input_file.endswith('.txt'):
            if output_format == 'pdf':
                with open(input_file, 'r', encoding='utf-8') as file:
                    text = file.read()
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(200, 10, txt=text)
                    pdf.output(output_path)
            elif output_format == 'docx':
                with open(input_file, 'r', encoding='utf-8') as file:
                    text = file.read()
                    doc = Document()
                    doc.add_paragraph(text)
                    doc.save(output_path)
            else:
                return None  # Same format conversion not supported

        elif input_file.endswith('.docx'):
            if output_format == 'txt':
                doc = Document(input_file)
                text = ''.join(para.text + '\n' for para in doc.paragraphs)
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(text)
            elif output_format == 'pdf':
                doc = Document(input_file)
                text = replace_problematic_characters(para.text + '\n' for para in doc.paragraphs)
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(200, 10, txt=text)
                pdf.output(output_path)
            else:
                return None  # Same format conversion not supported

        else:
            return None  # Unsupported file format

        return output_filename

    except Exception as e:
        print(f"Error during conversion: {e}")
        return None

def download_file(filename):
    if filename:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            return flask.send_from_directory(app.config['UPLOAD_FOLDER'],
                                              filename, as_attachment=True)
        except Exception as e:
            print(f"Error downloading file: {e}")
            return None

@app.route('/')
def index():
    return flask.render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    uploaded_file = flask.request.files['file']
    output_format = flask.request.form['output_format']

    if uploaded_file and allowed_file(uploaded_file.filename):
        uploaded_filename = upload_file(uploaded_file)
        if uploaded_filename:
            converted_filename = convert_file(os.path.join(app.config['UPLOAD_FOLDER'], uploaded_filename), output_format)
            if converted_filename:
                return download_file(converted_filename)
            else:
                flask.flash("Conversion failed. Please check the input file format and try again.", 'error')
        else:
            flask.flash("File upload failed.", 'error')
    else:
        flask.flash("Invalid file type. Please upload a TXT, PDF, or DOCX file.", 'error')

    return flask.redirect(flask.url_for('index'))

if __name__ == '__main__':
    app.config['UPLOAD_FOLDER'] = 'uploads'  # Set upload folder path
    try:
        os.makedirs(app.config['UPLOAD_FOLDER'])  # Create upload folder if it doesn't exist
    except OSError:
        pass
    app.run(debug=True)