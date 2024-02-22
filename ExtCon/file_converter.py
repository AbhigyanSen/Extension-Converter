import os
import PyPDF2
from docx import Document
from fpdf import FPDF

def replace_problematic_characters(text):
    problematic_chars = ['\u2019', '\u2013']  # Add other problematic characters if needed
    for char in problematic_chars:
        text = text.replace(char, '?')
    return text

def pdf_to_txt(input_path, output_dir):
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''.join(reader.pages[page_num].extract_text() for page_num in range(len(reader.pages)))

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.txt')
    with open(output_path, 'w') as txt_file:
        txt_file.write(text)

def pdf_to_docx(input_path, output_dir):
    with open(input_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        doc = Document()
        for page_num in range(len(reader.pages)):
            text = reader.pages[page_num].extract_text()
            doc.add_paragraph(text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.docx')
    doc.save(output_path)

def txt_to_pdf(input_path, output_dir):
    with open(input_path, 'r', encoding='utf-8') as file:
        text = replace_problematic_characters(file.read())

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, txt=text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.pdf')
    pdf.output(output_path)

def txt_to_docx(input_path, output_dir):
    with open(input_path, 'r') as file:
        text = file.read()

    doc = Document()
    doc.add_paragraph(text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.docx')
    doc.save(output_path)

def docx_to_pdf(input_path, output_dir):
    doc = Document(input_path)
    text = replace_problematic_characters(para.text + '\n' for para in doc.paragraphs)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, txt=text)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.pdf')
    pdf.output(output_path)

def docx_to_txt(input_path, output_dir):
    doc = Document(input_path)
    text = ''.join(para.text + '\n' for para in doc.paragraphs)

    output_path = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + '.txt')
    with open(output_path, 'w') as txt_file:
        txt_file.write(text)

path = input("Enter file: ")
ext = os.path.splitext(path)[-1].lower()
print(ext)

output_dir = "P:\\DataCore\\ExtCon\\uploads\\"
output = input("Enter output type (txt/pdf/docx): ")
output = "." + output.lower()
print(output)

if ext == ".docx":
    if output == ".txt":
        docx_to_txt(path, output_dir)
    elif output == ".pdf":
        docx_to_pdf(path, output_dir)
    else:
        print("SAME FILE TYPE CONVERSION DETECTED! NOT PROCEEDING!")
elif ext == ".txt":
    if output == ".pdf":
        txt_to_pdf(path, output_dir)
    elif output == ".docx":
        txt_to_docx(path, output_dir)
    else:
        print("SAME FILE TYPE CONVERSION DETECTED! NOT PROCEEDING!")
elif ext == ".pdf":
    if output == ".txt":
        pdf_to_txt(path, output_dir)
    elif output == ".docx":
        pdf_to_docx(path, output_dir)
    else:
        print("SAME FILE TYPE CONVERSION DETECTED! NOT PROCEEDING!")
else:
    print("INVALID FILE TYPE DETECTED! NOT PROCEEDING!")
